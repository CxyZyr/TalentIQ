"""
Word文档导出工具
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any
from datetime import datetime
import os


class WordExporter:
    """Word文档导出工具类"""

    def export_jd_to_word(self, jd: Any, output_path: str = None) -> str:
        """
        将JD导出为Word文档

        Args:
            jd: JobDescription对象
            output_path: 输出路径（可选）

        Returns:
            生成的文件路径
        """
        # 创建文档
        doc = Document()

        # 设置文档标题
        title = doc.add_heading(jd.job_title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加基本信息
        self._add_section(doc, "基本信息")
        basic_info = [
            ("岗位名称", jd.job_title),
            ("所属行业", jd.industry or "未填写"),
            ("岗位级别", jd.job_level or "未填写"),
            ("所属部门", jd.department_ref.name if jd.department_ref else jd.department),
            ("薪资范围", jd.salary_range or "面议"),
            ("岗位人数", str(jd.headcount) if jd.headcount else "未填写"),
            ("期望到岗时间", jd.expected_onboard_date.strftime("%Y-%m-%d") if jd.expected_onboard_date else "未填写"),
        ]

        table = doc.add_table(rows=len(basic_info), cols=2)
        table.style = 'Light Grid Accent 1'

        for i, (label, value) in enumerate(basic_info):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value

        # 添加岗位职责
        if jd.job_responsibilities:
            self._add_section(doc, "岗位职责")
            self._add_content(doc, jd.job_responsibilities)

        # 添加任职资格-硬性条件
        if jd.hard_requirements:
            self._add_section(doc, "任职资格 - 硬性条件")
            self._add_content(doc, jd.hard_requirements)

        # 添加任职资格-其他要求
        if jd.other_requirements:
            self._add_section(doc, "任职资格 - 其他要求")
            self._add_content(doc, jd.other_requirements)

        # 添加页脚信息
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.add_run(f"创建时间: {jd.created_at.strftime('%Y-%m-%d %H:%M:%S')}").font.size = Pt(9)
        footer.add_run(f"\n更新时间: {jd.updated_at.strftime('%Y-%m-%d %H:%M:%S')}").font.size = Pt(9)
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 生成文件路径
        if not output_path:
            # 创建输出目录
            output_dir = "exports"
            os.makedirs(output_dir, exist_ok=True)

            # 生成文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"JD_{jd.job_title}_{timestamp}.docx"
            output_path = os.path.join(output_dir, filename)

        # 保存文档
        doc.save(output_path)
        return output_path

    def _add_section(self, doc: Document, title: str):
        """添加章节标题"""
        heading = doc.add_heading(title, level=1)
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 139)

    def _add_content(self, doc: Document, content: str):
        """添加内容段落"""
        # 处理多行内容
        lines = content.split('\n')
        for line in lines:
            if line.strip():
                p = doc.add_paragraph(line.strip())
                p.paragraph_format.line_spacing = 1.5
